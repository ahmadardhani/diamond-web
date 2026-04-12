"""Tests covering remaining gaps in:
- diamond_web/utils/docx_template.py (line 70: continue when no {{row. placeholder)
- diamond_web/views/docx_template.py (30-31, 86-92, 182-187)
"""
import pytest
from unittest.mock import patch, MagicMock
from django.urls import reverse
from django.contrib.auth.models import Group
from django.http import HttpResponseRedirect

from diamond_web.tests.conftest import UserFactory, DocxTemplateFactory
from diamond_web.views.mixins import SafeDeleteMixin


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _admin_p3de_user():
    user = UserFactory()
    grp, _ = Group.objects.get_or_create(name='admin_p3de')
    user.groups.add(grp)
    return user


# ===========================================================================
# utils/docx_template.py
# ===========================================================================

@pytest.mark.django_db
class TestDocxTemplateUtilsGaps:
    """Cover remaining uncovered lines in utils/docx_template.py."""

    def test_fill_row_placeholders_no_placeholder_line_70(self):
        """Line 70 (`continue`): paragraph with no {{row. text is skipped."""
        from diamond_web.utils.docx_template import _fill_row_placeholders
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        # Cell 0: regular text with NO {{row. placeholder → hits line 70 (continue)
        table.rows[0].cells[0].paragraphs[0].add_run('Regular text, no placeholder')
        # Cell 1: text WITH {{row. placeholder → goes into replacement logic
        table.rows[0].cells[1].paragraphs[0].add_run('{{row.name}}')

        _fill_row_placeholders(table.rows[0], {'name': 'Alice'})

        # Cell 0 text unchanged (the continue skipped it)
        assert 'Regular text' in table.rows[0].cells[0].text
        # Cell 1 was replaced
        assert 'Alice' in table.rows[0].cells[1].text


# ===========================================================================
# views/docx_template.py
# ===========================================================================

@pytest.mark.django_db
class TestDocxTemplateViewListGaps:
    """Cover lines 30-31 in DocxTemplateListView.get (try block body)."""

    def test_list_view_with_deleted_and_name_covers_lines_30_31(self, client):
        """Lines 30-31: try block entered when deleted and name are both present."""
        user = _admin_p3de_user()
        client.force_login(user)

        # Both params present → enters the if block → covers lines 30-31 (try/unquote_plus)
        resp = client.get(
            reverse('docx_template_list'),
            {'deleted': 'true', 'name': 'My+Template+Name'},
        )
        assert resp.status_code == 200

    def test_list_view_unquote_plus_exception_covers_lines_33_34(self, client):
        """Lines 33-34: except Exception inside the try block.

        patch unquote_plus to raise so the except branch is taken.
        """
        user = _admin_p3de_user()
        client.force_login(user)

        with patch(
            'diamond_web.views.docx_template.unquote_plus',
            side_effect=Exception('mocked error'),
        ):
            resp = client.get(
                reverse('docx_template_list'),
                {'deleted': 'true', 'name': 'SomeName'},
            )
        assert resp.status_code == 200  # Exception caught with `pass`, view proceeds normally


@pytest.mark.django_db
class TestDocxTemplateDeleteViewGaps:
    """Cover lines 86-92 in DocxTemplateDeleteView.delete.

    In Django 5.x, BaseDeleteView.post uses form_valid() instead of calling
    self.delete() directly. Therefore DocxTemplateDeleteView.delete is NEVER
    invoked via normal HTTP dispatch. We call it directly via RequestFactory.
    """

    def test_delete_non_302_path_lines_86_89_92(self, db):
        """Lines 86-89, 92: delete body when super().delete() returns non-302."""
        from diamond_web.views.docx_template import DocxTemplateDeleteView
        from django.test import RequestFactory
        from django.http import JsonResponse as _JsonResponse

        user = _admin_p3de_user()
        obj = DocxTemplateFactory()

        factory = RequestFactory()
        request = factory.post(f'/docx-template/{obj.pk}/delete/')
        request.user = user

        view = DocxTemplateDeleteView()
        view.request = request
        view.kwargs = {'pk': obj.pk}
        view.args = ()

        with patch.object(
            SafeDeleteMixin,
            'delete',
            return_value=_JsonResponse({'success': True}),
        ):
            response = view.delete(request, pk=obj.pk)

        assert response.status_code == 200

    def test_delete_302_path_lines_90_91(self, db):
        """Lines 90-91: redirect branch when super().delete() returns 302."""
        from diamond_web.views.docx_template import DocxTemplateDeleteView
        from django.test import RequestFactory

        user = _admin_p3de_user()
        obj = DocxTemplateFactory(nama_template='Template Redirect Test')

        factory = RequestFactory()
        request = factory.post(f'/docx-template/{obj.pk}/delete/')
        request.user = user

        view = DocxTemplateDeleteView()
        view.request = request
        view.kwargs = {'pk': obj.pk}
        view.args = ()

        with patch.object(
            SafeDeleteMixin,
            'delete',
            return_value=HttpResponseRedirect('/'),
        ):
            response = view.delete(request, pk=obj.pk)

        assert response.status_code == 302
        assert 'deleted=true' in response.url


@pytest.mark.django_db
class TestDocxTemplateDownloadViewGaps:
    """Cover lines 182-187 in docx_template_download (exception handlers)."""

    def _make_admin(self):
        user = UserFactory()
        user.is_superuser = True
        user.save()
        return user

    def test_file_not_found_error_lines_182_184(self, client):
        """Lines 182-184: FileNotFoundError when opening template file → Http404."""
        user = self._make_admin()
        client.force_login(user)
        obj = DocxTemplateFactory()

        with patch.object(
            obj.file_template,
            'open',
            side_effect=FileNotFoundError('file missing'),
        ):
            # We need to patch on the model instance's field, but the view
            # re-fetches the object. Patch at a higher level:
            with patch(
                'django.db.models.fields.files.FieldFile.open',
                side_effect=FileNotFoundError('file missing'),
            ):
                resp = client.get(
                    reverse('docx_template_download', args=[obj.pk])
                )
        assert resp.status_code == 404

    def test_generic_exception_lines_185_187(self, client):
        """Lines 185-187: generic Exception when opening template file → Http404."""
        user = self._make_admin()
        client.force_login(user)
        obj = DocxTemplateFactory()

        with patch(
            'django.db.models.fields.files.FieldFile.open',
            side_effect=PermissionError('permission denied'),
        ):
            resp = client.get(
                reverse('docx_template_download', args=[obj.pk])
            )
        assert resp.status_code == 404
